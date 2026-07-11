"""
Comprehensive tests for agents: CRUD, documents, KB access, admin panel,
document processor, and agent RAG service.
"""
import io
import pytest
from unittest.mock import AsyncMock, MagicMock, patch

pytestmark = pytest.mark.asyncio

from tests.conftest import TEST_USER_ID, TEST_CONV_ID, make_token

TEST_AGENT_ID = "aaaaaaaa-aaaa-aaaa-aaaa-aaaaaaaaaaaa"
TEST_DOC_ID   = "bbbbbbbb-bbbb-bbbb-bbbb-bbbbbbbbbbbb"
ADMIN_EMAIL   = "katochaneesh@gmail.com"


def admin_token() -> str:
    import time, jwt, os
    secret = os.environ["SUPABASE_JWT_SECRET"]
    now = int(time.time())
    payload = {"sub": "admin-user-id", "email": ADMIN_EMAIL, "role": "authenticated", "iat": now, "exp": now + 3600}
    return jwt.encode(payload, secret, algorithm="HS256")


def _agent(agent_id=TEST_AGENT_ID, user_id=TEST_USER_ID):
    return {"id": agent_id, "user_id": user_id, "name": "Diet Agent", "description": "Health assistant",
            "system_prompt": "You are a diet expert.", "barge_in_sensitivity": 0.20,
            "created_at": "2026-06-28T00:00:00+00:00", "updated_at": "2026-06-28T00:00:00+00:00"}


def _doc(doc_id=TEST_DOC_ID, agent_id=TEST_AGENT_ID, scope="personal"):
    return {"id": doc_id, "agent_id": agent_id, "user_id": TEST_USER_ID, "file_name": "plan.txt",
            "file_size": 1024, "file_type": "txt", "scope": scope, "created_at": "2026-06-28T00:00:00+00:00"}


def _kb_req(status="pending"):
    return {"id": "req-1", "user_id": TEST_USER_ID, "user_email": "user@test.com",
            "status": status, "created_at": "2026-06-28T00:00:00+00:00",
            "resolved_at": None, "resolved_by": None}


@pytest.fixture
async def client():
    from httpx import AsyncClient, ASGITransport
    from app.main import app
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as ac:
        yield ac


def _auth(token=None):
    return {"Authorization": f"Bearer {token or make_token()}"}


def _admin_auth():
    return {"Authorization": f"Bearer {admin_token()}"}


# ─────────────────────────────────────────────────────────────────────────────
# Agent CRUD
# ─────────────────────────────────────────────────────────────────────────────

async def test_list_agents_empty(client, mocker):
    mocker.patch("app.agent_repo.get_agents", return_value=[])
    res = await client.get("/api/agents", headers=_auth())
    assert res.status_code == 200
    assert res.json() == []


async def test_list_agents_returns_agents(client, mocker):
    mocker.patch("app.agent_repo.get_agents", return_value=[_agent()])
    res = await client.get("/api/agents", headers=_auth())
    assert res.status_code == 200
    data = res.json()
    assert len(data) == 1
    assert data[0]["name"] == "Diet Agent"


async def test_list_agents_requires_auth(client):
    res = await client.get("/api/agents")
    assert res.status_code == 403


async def test_create_agent_success(client, mocker):
    mocker.patch("app.agent_repo.create_agent", return_value=_agent())
    res = await client.post("/api/agents", headers=_auth(), json={
        "name": "Diet Agent", "description": "Health assistant",
        "system_prompt": "You are a diet expert.", "barge_in_sensitivity": 0.20,
    })
    assert res.status_code == 201
    assert res.json()["name"] == "Diet Agent"


async def test_create_agent_minimal_fields(client, mocker):
    mocker.patch("app.agent_repo.create_agent", return_value=_agent())
    res = await client.post("/api/agents", headers=_auth(), json={"name": "Minimal Agent"})
    assert res.status_code == 201


async def test_create_agent_missing_name_returns_422(client, mocker):
    mocker.patch("app.agent_repo.create_agent", return_value=_agent())
    res = await client.post("/api/agents", headers=_auth(), json={"description": "No name agent"})
    assert res.status_code == 422


async def test_create_agent_requires_auth(client):
    res = await client.post("/api/agents", json={"name": "Test"})
    assert res.status_code == 403


async def test_get_agent_success(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    res = await client.get(f"/api/agents/{TEST_AGENT_ID}", headers=_auth())
    assert res.status_code == 200
    assert res.json()["id"] == TEST_AGENT_ID


async def test_get_agent_not_found(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=None)
    res = await client.get(f"/api/agents/{TEST_AGENT_ID}", headers=_auth())
    assert res.status_code == 404


async def test_get_agent_other_user_returns_404(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent(user_id="other-user"))
    res = await client.get(f"/api/agents/{TEST_AGENT_ID}", headers=_auth())
    assert res.status_code == 404


async def test_update_agent_success(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.update_agent", return_value=_agent())
    res = await client.put(f"/api/agents/{TEST_AGENT_ID}", headers=_auth(), json={
        "name": "Updated Diet Agent", "barge_in_sensitivity": 0.35,
    })
    assert res.status_code == 200


async def test_update_agent_not_found(client, mocker):
    # update route calls update_agent directly (no get_agent preflight); ownership
    # is enforced via eq("user_id") in SQL — RuntimeError → 404
    mocker.patch("app.agent_repo.update_agent", side_effect=RuntimeError("Agent not found or not owned by user"))
    res = await client.put(f"/api/agents/{TEST_AGENT_ID}", headers=_auth(), json={"name": "X"})
    assert res.status_code == 404


async def test_delete_agent_success(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.delete_agent", return_value=None)
    res = await client.delete(f"/api/agents/{TEST_AGENT_ID}", headers=_auth())
    assert res.status_code == 204


async def test_delete_agent_other_user_returns_404(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent(user_id="other-user"))
    res = await client.delete(f"/api/agents/{TEST_AGENT_ID}", headers=_auth())
    assert res.status_code == 404


async def test_delete_agent_requires_auth(client):
    res = await client.delete(f"/api/agents/{TEST_AGENT_ID}")
    assert res.status_code == 403


# ─────────────────────────────────────────────────────────────────────────────
# Document upload + listing
# ─────────────────────────────────────────────────────────────────────────────

async def test_list_documents_empty(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_agent_documents", return_value=[])
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)
    res = await client.get(f"/api/agents/{TEST_AGENT_ID}/documents", headers=_auth())
    assert res.status_code == 200
    data = res.json()
    assert data["documents"] == []
    assert data["storage_used_bytes"] == 0
    assert data["storage_limit_bytes"] == 5 * 1024 * 1024


async def test_list_documents_with_items(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_agent_documents", return_value=[_doc()])
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=1024)
    res = await client.get(f"/api/agents/{TEST_AGENT_ID}/documents", headers=_auth())
    assert res.status_code == 200
    data = res.json()
    assert len(data["documents"]) == 1
    assert data["storage_used_bytes"] == 1024


async def test_upload_txt_document_success(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)
    mocker.patch("app.agent_repo.create_agent_document", return_value=_doc())
    mocker.patch("app.agent_doc_processor.process_document", return_value=5)

    files = {"file": ("diet.txt", b"A diet plan with lots of tips for healthy eating.", "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 201
    assert res.json()["file_name"] == "plan.txt"


async def test_upload_pdf_document_success(client, mocker):
    import fitz
    fitz_doc = fitz.open()
    page = fitz_doc.new_page()
    page.insert_text((50, 100), "PDF diet content.")
    buf = io.BytesIO()
    fitz_doc.save(buf)

    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)
    mocker.patch("app.agent_repo.create_agent_document", return_value=_doc(doc_id="pdf-doc"))
    mocker.patch("app.agent_doc_processor.process_document", return_value=3)

    files = {"file": ("diet.pdf", buf.getvalue(), "application/pdf")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 201


async def test_upload_unsupported_format_rejected(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)

    files = {"file": ("notes.csv", b"col1,col2\nval1,val2", "text/csv")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 400
    assert "PDF" in res.json()["detail"]


async def test_upload_over_5mb_limit_rejected(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=5 * 1024 * 1024)

    files = {"file": ("extra.txt", b"x" * 100, "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 413
    assert "storage" in res.json()["detail"].lower() or "limit" in res.json()["detail"].lower()


async def test_upload_near_limit_allowed(client, mocker):
    """File that fits exactly within remaining space should succeed."""
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=5 * 1024 * 1024 - 50)
    mocker.patch("app.agent_repo.create_agent_document", return_value=_doc())
    mocker.patch("app.agent_doc_processor.process_document", return_value=1)

    content = b"A " * 25  # 50 bytes, exactly fits remaining space
    files = {"file": ("tiny.txt", content, "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 201


async def test_upload_global_without_permission_rejected(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)

    files = {"file": ("guide.txt", b"Some knowledge base content.", "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "global"})
    assert res.status_code == 403
    assert "admin approval" in res.json()["detail"].lower()


async def test_upload_global_with_permission_succeeds(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=True)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)
    mocker.patch("app.agent_repo.create_agent_document", return_value=_doc(scope="global"))
    mocker.patch("app.agent_doc_processor.process_document", return_value=2)

    files = {"file": ("global.txt", b"Global knowledge base content here.", "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "global"})
    assert res.status_code == 201
    assert res.json()["scope"] == "global"


async def test_upload_empty_document_rejected(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)

    files = {"file": ("empty.txt", b"", "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 422


async def test_upload_whitespace_only_document_rejected(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    mocker.patch("app.agent_repo.get_agent_storage_used", return_value=0)

    files = {"file": ("blank.txt", b"   \n\t  \n  ", "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 422


async def test_delete_document_success(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent())
    mocker.patch("app.agent_repo.delete_agent_document", return_value=None)
    res = await client.delete(f"/api/agents/{TEST_AGENT_ID}/documents/{TEST_DOC_ID}",
                               headers=_auth())
    assert res.status_code == 204


async def test_upload_requires_agent_ownership(client, mocker):
    mocker.patch("app.agent_repo.get_agent", return_value=_agent(user_id="other-user"))
    files = {"file": ("test.txt", b"content", "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            headers=_auth(), files=files, data={"scope": "personal"})
    assert res.status_code == 404


async def test_upload_requires_auth(client):
    files = {"file": ("test.txt", b"content", "text/plain")}
    res = await client.post(f"/api/agents/{TEST_AGENT_ID}/documents",
                            files=files, data={"scope": "personal"})
    assert res.status_code == 403


# ─────────────────────────────────────────────────────────────────────────────
# KB Access Requests
# ─────────────────────────────────────────────────────────────────────────────

async def test_my_kb_permission_false(client, mocker):
    mocker.patch("app.agent_repo.get_kb_permission", return_value=False)
    res = await client.get("/api/kb/my-permission", headers=_auth())
    assert res.status_code == 200
    assert res.json()["can_upload_global_kb"] is False


async def test_my_kb_permission_true(client, mocker):
    mocker.patch("app.agent_repo.get_kb_permission", return_value=True)
    res = await client.get("/api/kb/my-permission", headers=_auth())
    assert res.status_code == 200
    assert res.json()["can_upload_global_kb"] is True


async def test_request_kb_access_success(client, mocker):
    mocker.patch("app.agent_repo.create_kb_request", return_value=_kb_req())
    res = await client.post("/api/kb/request-access", headers=_auth(), json={"user_email": "user@test.com"})
    assert res.status_code == 200
    assert res.json()["status"] == "pending"


async def test_kb_request_requires_auth(client):
    res = await client.post("/api/kb/request-access", json={"user_email": "x@x.com"})
    assert res.status_code == 403


async def test_kb_request_missing_email_returns_422(client):
    res = await client.post("/api/kb/request-access", headers=_auth(), json={})
    assert res.status_code == 422


# ─────────────────────────────────────────────────────────────────────────────
# Admin Routes
# ─────────────────────────────────────────────────────────────────────────────

async def test_is_admin_false_for_regular_user(client):
    res = await client.get("/api/admin/is-admin", headers=_auth())
    assert res.status_code == 200
    assert res.json()["is_admin"] is False


async def test_is_admin_true_for_admin_email(client):
    res = await client.get("/api/admin/is-admin", headers=_admin_auth())
    assert res.status_code == 200
    assert res.json()["is_admin"] is True


async def test_is_admin_requires_auth(client):
    res = await client.get("/api/admin/is-admin")
    assert res.status_code == 403


async def test_admin_list_requests_forbidden_for_non_admin(client):
    res = await client.get("/api/admin/kb-requests", headers=_auth())
    assert res.status_code == 403


async def test_admin_list_requests_success(client, mocker):
    mocker.patch("app.agent_repo.get_kb_requests", return_value=[_kb_req()])
    res = await client.get("/api/admin/kb-requests", headers=_admin_auth())
    assert res.status_code == 200
    data = res.json()
    assert isinstance(data, list)
    assert data[0]["status"] == "pending"


async def test_admin_list_requests_empty(client, mocker):
    mocker.patch("app.agent_repo.get_kb_requests", return_value=[])
    res = await client.get("/api/admin/kb-requests", headers=_admin_auth())
    assert res.status_code == 200
    assert res.json() == []


async def test_admin_approve_request_success(client, mocker):
    mocker.patch("app.agent_repo.resolve_kb_request", return_value=_kb_req("approved"))
    res = await client.post("/api/admin/kb-requests/req-1/approve", headers=_admin_auth())
    assert res.status_code == 200
    assert res.json()["status"] == "approved"


async def test_admin_reject_request_success(client, mocker):
    mocker.patch("app.agent_repo.resolve_kb_request", return_value=_kb_req("rejected"))
    res = await client.post("/api/admin/kb-requests/req-1/reject", headers=_admin_auth())
    assert res.status_code == 200
    assert res.json()["status"] == "rejected"


async def test_admin_approve_request_not_found(client, mocker):
    mocker.patch("app.agent_repo.resolve_kb_request", side_effect=RuntimeError("Request not found"))
    res = await client.post("/api/admin/kb-requests/bad-req/approve", headers=_admin_auth())
    assert res.status_code == 404


async def test_admin_approve_forbidden_for_non_admin(client):
    res = await client.post("/api/admin/kb-requests/req-1/approve", headers=_auth())
    assert res.status_code == 403


async def test_admin_reject_forbidden_for_non_admin(client):
    res = await client.post("/api/admin/kb-requests/req-1/reject", headers=_auth())
    assert res.status_code == 403


async def test_admin_revoke_kb_success(client, mocker):
    mocker.patch("app.agent_repo.revoke_kb_permission", return_value=None)
    res = await client.post(f"/api/admin/users/{TEST_USER_ID}/revoke-kb", headers=_admin_auth())
    assert res.status_code == 200
    assert res.json()["status"] == "revoked"


async def test_admin_revoke_kb_forbidden_for_non_admin(client):
    res = await client.post(f"/api/admin/users/{TEST_USER_ID}/revoke-kb", headers=_auth())
    assert res.status_code == 403


async def test_admin_list_kb_users_success(client, mocker):
    users = [{"user_id": TEST_USER_ID, "can_upload_global_kb": True, "granted_at": None, "granted_by": None}]
    mocker.patch("app.agent_repo.get_kb_permitted_users", return_value=users)
    res = await client.get("/api/admin/kb-users", headers=_admin_auth())
    assert res.status_code == 200
    assert isinstance(res.json(), list)
    assert res.json()[0]["user_id"] == TEST_USER_ID


async def test_admin_list_kb_users_forbidden_for_non_admin(client):
    res = await client.get("/api/admin/kb-users", headers=_auth())
    assert res.status_code == 403


# ─────────────────────────────────────────────────────────────────────────────
# Document Processor Unit Tests
# ─────────────────────────────────────────────────────────────────────────────

def test_extract_text_txt_basic():
    from app.agent_doc_processor import extract_text_txt
    result = extract_text_txt(b"Hello, this is a text file.\nSecond line.")
    assert "Hello" in result
    assert "Second line" in result


def test_extract_text_txt_unicode():
    from app.agent_doc_processor import extract_text_txt
    result = extract_text_txt("Héllo Wörld — diet plan.".encode("utf-8"))
    assert "diet plan" in result


def test_extract_text_txt_bad_bytes_uses_replacement():
    from app.agent_doc_processor import extract_text_txt
    result = extract_text_txt(b"hello \xff\xfe world")
    assert "hello" in result
    assert "world" in result


def test_extract_text_docx():
    from app.agent_doc_processor import extract_text_docx
    from docx import Document
    doc = Document()
    doc.add_paragraph("Diet plan paragraph one.")
    doc.add_paragraph("Eat more vegetables daily.")
    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text_docx(buf.getvalue())
    assert "Diet plan paragraph one." in result
    assert "vegetables" in result


def test_extract_text_docx_skips_whitespace_paragraphs():
    from app.agent_doc_processor import extract_text_docx
    from docx import Document
    doc = Document()
    doc.add_paragraph("   ")  # whitespace-only → skipped
    doc.add_paragraph("Real content here.")
    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text_docx(buf.getvalue())
    assert "Real content here." in result


def test_extract_text_pdf():
    from app.agent_doc_processor import extract_text_pdf
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), "PDF diet plan content here.")
    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text_pdf(buf.getvalue())
    assert isinstance(result, str)
    assert len(result) > 0


def test_extract_text_dispatch_pdf():
    from app.agent_doc_processor import extract_text
    import fitz
    doc = fitz.open()
    doc.new_page()
    buf = io.BytesIO()
    doc.save(buf)
    assert isinstance(extract_text(buf.getvalue(), "pdf"), str)


def test_extract_text_dispatch_txt():
    from app.agent_doc_processor import extract_text
    assert extract_text(b"hello world", "txt") == "hello world"


def test_extract_text_unsupported_raises():
    from app.agent_doc_processor import extract_text
    with pytest.raises(ValueError, match="Unsupported"):
        extract_text(b"data", "csv")


def test_extract_text_unsupported_md_raises():
    from app.agent_doc_processor import extract_text
    with pytest.raises(ValueError, match="Unsupported"):
        extract_text(b"# heading", "md")


def test_chunk_text_short_stays_one_chunk():
    from app.agent_doc_processor import chunk_text
    text = "This is a short text with only a few words."
    chunks = chunk_text(text)
    assert len(chunks) == 1
    assert chunks[0] == text


def test_chunk_text_empty_returns_empty():
    from app.agent_doc_processor import chunk_text
    assert chunk_text("") == []
    assert chunk_text("   ") == []


def test_chunk_text_splits_large_text():
    from app.agent_doc_processor import chunk_text
    text = " ".join([f"word{i}" for i in range(1000)])
    chunks = chunk_text(text, chunk_words=100, overlap=10)
    assert len(chunks) > 1
    for c in chunks:
        assert len(c.split()) <= 105


def test_chunk_text_overlap_creates_continuity():
    from app.agent_doc_processor import chunk_text
    words = [f"w{i}" for i in range(50)]
    chunks = chunk_text(" ".join(words), chunk_words=20, overlap=5)
    if len(chunks) >= 2:
        # chunk1 = words[0:20], chunk2 starts at word 15 (20-5=15)
        assert chunks[1].split()[0] == "w15"


def test_chunk_text_covers_last_word():
    from app.agent_doc_processor import chunk_text
    words = [f"w{i}" for i in range(100)]
    chunks = chunk_text(" ".join(words), chunk_words=30, overlap=5)
    assert "w99" in chunks[-1].split()


def test_chunk_text_exactly_at_limit():
    from app.agent_doc_processor import chunk_text
    text = " ".join([f"w{i}" for i in range(400)])
    chunks = chunk_text(text, chunk_words=400, overlap=40)
    assert len(chunks) == 1


async def test_embed_chunks_calls_openai(mocker):
    from app.agent_doc_processor import embed_chunks
    mock_client = MagicMock()
    mock_client.embeddings.create.return_value.data = [
        MagicMock(embedding=[0.1] * 1536),
        MagicMock(embedding=[0.2] * 1536),
    ]
    mocker.patch("app.agent_doc_processor._openai", mock_client)

    result = embed_chunks(["chunk one", "chunk two"])
    assert len(result) == 2
    assert len(result[0]) == 1536
    assert result[1][0] == pytest.approx(0.2)
    mock_client.embeddings.create.assert_called_once()


async def test_embed_chunks_empty_returns_empty():
    from app.agent_doc_processor import embed_chunks
    result = embed_chunks([])
    assert result == []


async def test_store_chunks_inserts_correct_rows(mocker):
    from app.agent_doc_processor import store_chunks
    mock_db = MagicMock()
    mock_db.table.return_value.insert.return_value.execute.return_value.data = [{"id": "c1"}]
    mocker.patch("app.agent_doc_processor.get_supabase", return_value=mock_db)

    chunks = ["chunk 1", "chunk 2", "chunk 3"]
    embeddings = [[0.1] * 1536, [0.2] * 1536, [0.3] * 1536]
    count = store_chunks("doc-id", "agent-id", "user-id", "personal", chunks, embeddings)
    assert count == 3
    mock_db.table.assert_called_with("document_chunks")

    inserted_rows = mock_db.table.return_value.insert.call_args[0][0]
    assert len(inserted_rows) == 3
    assert inserted_rows[0]["chunk_index"] == 0
    assert inserted_rows[1]["chunk_index"] == 1
    assert inserted_rows[2]["chunk_index"] == 2
    assert inserted_rows[0]["scope"] == "personal"
    assert inserted_rows[0]["agent_id"] == "agent-id"
    assert inserted_rows[0]["document_id"] == "doc-id"


async def test_store_chunks_empty_does_not_insert(mocker):
    from app.agent_doc_processor import store_chunks
    mock_db = MagicMock()
    mocker.patch("app.agent_doc_processor.get_supabase", return_value=mock_db)

    count = store_chunks("doc-id", "agent-id", "user-id", "personal", [], [])
    assert count == 0
    mock_db.table.return_value.insert.assert_not_called()


async def test_process_document_full_pipeline(mocker):
    from app.agent_doc_processor import process_document
    mocker.patch("app.agent_doc_processor.embed_chunks", return_value=[[0.1] * 1536] * 3)
    mocker.patch("app.agent_doc_processor.store_chunks", return_value=3)

    content = " ".join([f"word{i}" for i in range(50)])
    count = process_document("doc-id", "agent-id", "user-id", "personal", content)
    assert count == 3


async def test_process_document_empty_content_returns_zero(mocker):
    from app.agent_doc_processor import process_document
    mock_embed = mocker.patch("app.agent_doc_processor.embed_chunks")
    count = process_document("doc-id", "agent-id", "user-id", "personal", "")
    assert count == 0
    mock_embed.assert_not_called()


# ─────────────────────────────────────────────────────────────────────────────
# Agent RAG Service Unit Tests
# ─────────────────────────────────────────────────────────────────────────────

def test_format_agent_context_empty():
    from app.agent_rag_service import format_agent_context
    assert format_agent_context([]) == ""


def test_format_agent_context_personal():
    from app.agent_rag_service import format_agent_context
    chunks = [{"content": "Eat less sugar.", "scope": "personal", "similarity": 0.9}]
    result = format_agent_context(chunks)
    assert "PERSONAL" in result
    assert "Eat less sugar." in result


def test_format_agent_context_global():
    from app.agent_rag_service import format_agent_context
    chunks = [{"content": "Global tip.", "scope": "global", "similarity": 0.95}]
    result = format_agent_context(chunks)
    assert "GLOBAL" in result
    assert "Global tip." in result


def test_format_agent_context_multiple_with_separator():
    from app.agent_rag_service import format_agent_context
    chunks = [
        {"content": "Global tip.", "scope": "global", "similarity": 0.95},
        {"content": "Personal note.", "scope": "personal", "similarity": 0.85},
    ]
    result = format_agent_context(chunks)
    assert "GLOBAL" in result
    assert "PERSONAL" in result
    assert "---" in result


def test_retrieve_agent_context_empty_scopes_returns_empty():
    from app.agent_rag_service import retrieve_agent_context
    result = retrieve_agent_context("query", "agent-id", "user-id", scopes=[])
    assert result == []


async def test_retrieve_agent_context_personal_scope(mocker):
    from app.agent_rag_service import retrieve_agent_context

    mock_openai = MagicMock()
    mock_openai.embeddings.create.return_value.data = [MagicMock(embedding=[0.1] * 1536)]
    mocker.patch("app.agent_rag_service._openai", mock_openai)

    mock_cursor = MagicMock()
    mock_cursor.__enter__ = MagicMock(return_value=mock_cursor)
    mock_cursor.__exit__ = MagicMock(return_value=False)
    mock_cursor.fetchall.return_value = [("Eat vegetables.", "personal", 0.92)]

    mock_conn = MagicMock()
    mock_conn.cursor.return_value = mock_cursor
    mocker.patch("app.agent_rag_service._get_conn", return_value=mock_conn)

    results = retrieve_agent_context("diet tips", "agent-id", "user-id", ["personal"])
    assert len(results) == 1
    assert results[0]["content"] == "Eat vegetables."
    assert results[0]["scope"] == "personal"
    assert results[0]["similarity"] == pytest.approx(0.92)


async def test_retrieve_agent_context_global_scope_no_user_id_filter(mocker):
    from app.agent_rag_service import retrieve_agent_context

    mock_openai = MagicMock()
    mock_openai.embeddings.create.return_value.data = [MagicMock(embedding=[0.2] * 1536)]
    mocker.patch("app.agent_rag_service._openai", mock_openai)

    mock_cursor = MagicMock()
    mock_cursor.__enter__ = MagicMock(return_value=mock_cursor)
    mock_cursor.__exit__ = MagicMock(return_value=False)
    mock_cursor.fetchall.return_value = [("Global health tip.", "global", 0.88)]

    mock_conn = MagicMock()
    mock_conn.cursor.return_value = mock_cursor
    mocker.patch("app.agent_rag_service._get_conn", return_value=mock_conn)

    results = retrieve_agent_context("health tips", "agent-id", "user-id", ["global"])
    assert results[0]["scope"] == "global"

    # Global scope should NOT filter by user_id
    called_sql = mock_cursor.execute.call_args[0][0]
    called_params = mock_cursor.execute.call_args[0][1]
    assert "scope = 'global'" in called_sql
    assert "user-id" not in called_params  # user_id not in params for global-only


async def test_retrieve_agent_context_both_scopes_includes_user_id_in_query(mocker):
    from app.agent_rag_service import retrieve_agent_context

    mock_openai = MagicMock()
    mock_openai.embeddings.create.return_value.data = [MagicMock(embedding=[0.3] * 1536)]
    mocker.patch("app.agent_rag_service._openai", mock_openai)

    mock_cursor = MagicMock()
    mock_cursor.__enter__ = MagicMock(return_value=mock_cursor)
    mock_cursor.__exit__ = MagicMock(return_value=False)
    mock_cursor.fetchall.return_value = []

    mock_conn = MagicMock()
    mock_conn.cursor.return_value = mock_cursor
    mocker.patch("app.agent_rag_service._get_conn", return_value=mock_conn)

    retrieve_agent_context("tips", "agent-id", "user-id", ["personal", "global"])

    called_sql = mock_cursor.execute.call_args[0][0]
    called_params = mock_cursor.execute.call_args[0][1]
    assert "user_id" in called_sql
    assert "scope = 'global'" in called_sql
    assert "user-id" in called_params


async def test_retrieve_agent_context_closes_connection_on_error(mocker):
    from app.agent_rag_service import retrieve_agent_context

    mock_openai = MagicMock()
    mock_openai.embeddings.create.return_value.data = [MagicMock(embedding=[0.1] * 1536)]
    mocker.patch("app.agent_rag_service._openai", mock_openai)

    mock_cursor = MagicMock()
    mock_cursor.__enter__ = MagicMock(return_value=mock_cursor)
    mock_cursor.__exit__ = MagicMock(return_value=False)
    mock_cursor.execute.side_effect = Exception("DB connection failed")

    mock_conn = MagicMock()
    mock_conn.cursor.return_value = mock_cursor
    mocker.patch("app.agent_rag_service._get_conn", return_value=mock_conn)

    with pytest.raises(Exception, match="DB connection failed"):
        retrieve_agent_context("query", "agent-id", "user-id", ["personal"])

    mock_conn.close.assert_called_once()


# ─────────────────────────────────────────────────────────────────────────────
# LLM system_prompt_override
# ─────────────────────────────────────────────────────────────────────────────

async def test_generate_response_stream_uses_custom_prompt(mocker):
    from app.llm_service import generate_response_stream

    captured_messages = []

    async def mock_create(**kwargs):
        captured_messages.extend(kwargs["messages"])
        async def _gen():
            c = MagicMock()
            c.choices[0].delta.content = "Custom reply."
            yield c
            c2 = MagicMock()
            c2.choices[0].delta.content = None
            yield c2
        return _gen()

    mock_client = AsyncMock()
    mock_client.chat.completions.create = mock_create
    mocker.patch("app.llm_service._get_openai", return_value=mock_client)

    async for _ in generate_response_stream(
        "What should I eat?", [], "", None,
        system_prompt_override="You are a strict diet coach. Only discuss food."
    ):
        pass

    system_msg = next((m for m in captured_messages if m["role"] == "system"), None)
    assert system_msg is not None
    assert "strict diet coach" in system_msg["content"]
    assert "voice AI assistant" not in system_msg["content"]


async def test_generate_response_stream_uses_default_when_no_override(mocker):
    from app.llm_service import generate_response_stream

    captured_messages = []

    async def mock_create(**kwargs):
        captured_messages.extend(kwargs["messages"])
        async def _gen():
            c = MagicMock()
            c.choices[0].delta.content = "Default reply."
            yield c
            c2 = MagicMock()
            c2.choices[0].delta.content = None
            yield c2
        return _gen()

    mock_client = AsyncMock()
    mock_client.chat.completions.create = mock_create
    mocker.patch("app.llm_service._get_openai", return_value=mock_client)

    async for _ in generate_response_stream("Hello", [], "", None):
        pass

    system_msg = next((m for m in captured_messages if m["role"] == "system"), None)
    assert system_msg is not None
    assert "voice AI assistant" in system_msg["content"]


# ─────────────────────────────────────────────────────────────────────────────
# Agent Repo Unit Tests
# ─────────────────────────────────────────────────────────────────────────────

def test_get_agent_storage_used_zero(mocker):
    from app.agent_repo import get_agent_storage_used
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.execute.return_value.data = []
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    assert get_agent_storage_used("agent-id") == 0


def test_get_agent_storage_used_sums_multiple_files(mocker):
    from app.agent_repo import get_agent_storage_used
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.execute.return_value.data = [
        {"file_size": 1024}, {"file_size": 2048}, {"file_size": 512},
    ]
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    assert get_agent_storage_used("agent-id") == 3584


def test_get_agent_storage_used_single_at_max(mocker):
    from app.agent_repo import get_agent_storage_used
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.execute.return_value.data = [
        {"file_size": 5 * 1024 * 1024},
    ]
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    assert get_agent_storage_used("agent-id") == 5 * 1024 * 1024


def test_get_kb_permission_no_row_returns_false(mocker):
    from app.agent_repo import get_kb_permission
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.limit.return_value.execute.return_value.data = []
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    assert get_kb_permission("user-id") is False


def test_get_kb_permission_true_row(mocker):
    from app.agent_repo import get_kb_permission
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.limit.return_value.execute.return_value.data = [
        {"can_upload_global_kb": True}
    ]
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    assert get_kb_permission("user-id") is True


def test_get_kb_permission_false_row(mocker):
    from app.agent_repo import get_kb_permission
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.limit.return_value.execute.return_value.data = [
        {"can_upload_global_kb": False}
    ]
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    assert get_kb_permission("user-id") is False


def test_create_kb_request_deduplicates_pending(mocker):
    from app.agent_repo import create_kb_request
    existing = [{"id": "req-1", "status": "pending"}]
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.eq.return_value.limit.return_value.execute.return_value.data = existing
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    result = create_kb_request("user-id", "user@email.com")
    assert result["id"] == "req-1"
    mock_db.table.return_value.insert.assert_not_called()


def test_create_kb_request_creates_new_when_no_pending(mocker):
    from app.agent_repo import create_kb_request
    mock_db = MagicMock()
    mock_db.table.return_value.select.return_value.eq.return_value.eq.return_value.limit.return_value.execute.return_value.data = []
    mock_db.table.return_value.insert.return_value.execute.return_value.data = [{"id": "new-req", "status": "pending"}]
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    result = create_kb_request("user-id", "user@email.com")
    assert result["id"] == "new-req"
    mock_db.table.return_value.insert.assert_called_once()


def test_create_agent_db_failure_raises_runtimeerror(mocker):
    from app.agent_repo import create_agent
    mock_db = MagicMock()
    mock_db.table.return_value.insert.return_value.execute.return_value.data = []
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    with pytest.raises(RuntimeError, match="Failed to create agent"):
        create_agent("user-id", "Test Agent")


def test_update_agent_db_failure_raises_runtimeerror(mocker):
    from app.agent_repo import update_agent
    mock_db = MagicMock()
    mock_db.table.return_value.update.return_value.eq.return_value.eq.return_value.execute.return_value.data = []
    mocker.patch("app.agent_repo.get_supabase", return_value=mock_db)
    with pytest.raises(RuntimeError, match="Agent not found"):
        update_agent("agent-id", "user-id", name="New Name")
